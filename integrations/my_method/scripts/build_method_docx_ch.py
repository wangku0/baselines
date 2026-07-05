from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("all-CH.docx").resolve()


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def shade_paragraph(paragraph, fill: str = "F4F6F9") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int]) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths_dxa[idx])


def add_field(paragraph, field: str) -> None:
    run = paragraph.add_run()
    fld_char_1 = OxmlElement("w:fldChar")
    fld_char_1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = field
    fld_char_2 = OxmlElement("w:fldChar")
    fld_char_2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_1)
    run._r.append(instr_text)
    run._r.append(fld_char_2)


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor(0, 0, 0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    title = styles["Title"]
    title.font.name = "Calibri"
    title._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    title.font.size = Pt(22)
    title.font.color.rgb = RGBColor(11, 37, 69)
    title.paragraph_format.space_after = Pt(6)

    subtitle = styles["Subtitle"]
    subtitle.font.name = "Calibri"
    subtitle._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    subtitle.font.size = Pt(11)
    subtitle.font.color.rgb = RGBColor(85, 85, 85)

    for name, size, color, before, after in [
        ("Heading 1", 16, RGBColor(46, 116, 181), 18, 10),
        ("Heading 2", 13, RGBColor(46, 116, 181), 12, 6),
        ("Heading 3", 12, RGBColor(31, 77, 120), 8, 4),
    ]:
        st = styles[name]
        st.font.name = "Calibri"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        st.font.size = Pt(size)
        st.font.color.rgb = color
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True

    code = styles.add_style("CodeBlock", 1)
    code.font.name = "Consolas"
    code._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    code.font.size = Pt(9)
    code.font.color.rgb = RGBColor(31, 58, 95)
    code.paragraph_format.left_indent = Inches(0.18)
    code.paragraph_format.right_indent = Inches(0.18)
    code.paragraph_format.space_before = Pt(3)
    code.paragraph_format.space_after = Pt(8)
    code.paragraph_format.line_spacing = 1.08

    formula = styles.add_style("Formula", 1)
    formula.font.name = "Cambria Math"
    formula._element.rPr.rFonts.set(qn("w:eastAsia"), "Cambria Math")
    formula.font.size = Pt(10.5)
    formula.font.color.rgb = RGBColor(0, 0, 0)
    formula.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    formula.paragraph_format.space_before = Pt(3)
    formula.paragraph_format.space_after = Pt(3)

    note = styles.add_style("NoteText", 1)
    note.font.name = "Calibri"
    note._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    note.font.size = Pt(10)
    note.font.color.rgb = RGBColor(31, 58, 95)
    note.paragraph_format.left_indent = Inches(0.16)
    note.paragraph_format.right_indent = Inches(0.16)
    note.paragraph_format.space_before = Pt(4)
    note.paragraph_format.space_after = Pt(8)
    note.paragraph_format.line_spacing = 1.25

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.add_run("风险子空间 Flow-guided LoRA 方法 | 第 ")
    add_field(footer, "PAGE")
    footer.add_run(" 页")


def add_para(doc: Document, text: str, style: str | None = None) -> None:
    doc.add_paragraph(text, style=style)


def add_formula(doc: Document, display: str, latex: str) -> None:
    p = doc.add_paragraph(display, style="Formula")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    label = doc.add_paragraph("LaTeX 代码：", style="NoteText")
    label.runs[0].bold = True
    code = doc.add_paragraph(latex.strip(), style="CodeBlock")
    shade_paragraph(code)


def add_note(doc: Document, title: str, text: str) -> None:
    p = doc.add_paragraph(style="NoteText")
    shade_paragraph(p)
    r = p.add_run(title + " ")
    r.bold = True
    p.add_run(text)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.375)
        p.paragraph_format.first_line_indent = Inches(-0.194)
        p.paragraph_format.space_after = Pt(4)
        p.add_run(item)


def add_summary_table(doc: Document) -> None:
    doc.add_heading("方法组成与作用概览", level=2)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    widths = [1900, 3950, 3510]
    set_table_geometry(table, widths)
    hdr = table.rows[0].cells
    for i, text in enumerate(["模块", "在方法中的作用", "这样设计的好处"]):
        hdr[i].text = text
        set_cell_shading(hdr[i], "F4F6F9")
        for p in hdr[i].paragraphs:
            for run in p.runs:
                run.bold = True
    rows = [
        (
            "风险子空间",
            "从 harmful 与 safe-neighbor 的成对隐状态差分中提取风险相关方向。",
            "抵消共享场景语义，突出 harmful 样本相对于安全邻居多出的危险操作偏移。",
        ),
        (
            "风险迁移选层",
            "选择 safe-to-harmful 迁移最强且 retain 重叠较低的编辑层。",
            "把 LoRA 修改限制在少数高影响层，降低对通用能力的副作用。",
        ),
        (
            "Flow Matching teacher",
            "学习 harmful 隐状态到 safe-neighbor 隐状态的连续迁移路径。",
            "提供平滑的潜空间安全迁移方向，避免只依赖文本监督造成模板化拒答。",
        ),
        (
            "Flow-guided LoRA",
            "结合 safe CE、风险投影约束、KL 保持、retain hidden 保持和 flow 蒸馏训练 LoRA。",
            "在冻结基础模型的前提下降低风险，同时尽量保持原模型能力。",
        ),
        (
            "统一评估",
            "报告显性 ASR、隐性风险子空间激活、融合总风险、拒答率、ROUGE 和清除率。",
            "同时观察可见有害行为、潜在风险激活和能力保持，避免单一指标误导。",
        ),
    ]
    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            cells[i].text = text
            set_cell_width(cells[i], widths[i])
    doc.add_paragraph()


def main() -> None:
    doc = Document()
    style_document(doc)

    doc.add_paragraph("风险子空间引导的 Flow-guided LoRA 安全遗忘方法", style="Title")
    doc.add_paragraph("包含详细原理、公式与可复用 LaTeX 代码的中文论文方法章节草稿", style="Subtitle")
    add_note(
        doc,
        "范围说明。",
        "本文档描述当前主方法：基于 SafeNb 的反事实风险子空间构建、risk-transport top-n 选层、Flow Matching teacher、Flow-guided LoRA unlearning，以及统一严格 before/after 评估。本文档不包含此前已撤回的 harmful-only 变体。",
    )

    add_summary_table(doc)

    doc.add_heading("1. 问题设定与数据组织", level=1)
    add_para(
        doc,
        "本文目标是在多模态大模型中降低 harmful trigger 诱导模型输出危险操作性回答的倾向，同时尽量保持 safe-neighbor 与 retain 样本上的正常能力。设基础多模态模型为 f_theta，输入由图像 v 与文本指令 x 共同构成。数据被组织为 harmful trigger、safe neighbor 与 retain 三类样本。",
    )
    add_formula(
        doc,
        r"D = {(x_i^h, x_i^s, x_i^r)}_{i=1}^N",
        r"""
\[
\mathcal{D}=\{(x_i^h, x_i^s, x_i^r)\}_{i=1}^{N}
\]
""",
    )
    add_para(
        doc,
        "其中，x_i^h 表示 harmful trigger，x_i^s 表示与其配对的 safe-neighbor 样本，x_i^r 表示 retain 样本。safe neighbor 并不是普通安全样本，而是在图像场景和任务语境上与 harmful trigger 高度接近、但回答目标为安全规避或非操作性处理的邻近样本。",
    )
    add_note(
        doc,
        "为什么这样做。",
        "如果直接用 harmful 与无关 retain 样本作对比，风险方向可能主要捕获主题、图像或场景分布差异。使用 paired safe-neighbor 作为反事实对照，可以在很大程度上控制视觉场景和任务模板，从而更集中地提取 unsafe operational deviation。",
    )

    doc.add_heading("2. 隐状态表示提取", level=1)
    add_para(
        doc,
        "对于每个多模态输入，方法从模型指定层提取隐状态。在当前实现中，使用第 l 层最后一个输入 token 的隐状态作为样本表示。该 token 位于模型开始生成回答之前，已经聚合了图像内容与完整文本指令，因此可以近似视为模型生成前的条件状态。",
    )
    add_formula(
        doc,
        r"h_l(x) = LastTokenHidden_l(f_theta(v,x))",
        r"""
\[
h_l(x)=\mathrm{LastTokenHidden}_l(f_\theta(v,x))
\]
""",
    )
    add_para(doc, "其中 l 属于候选层集合 L，例如 {4, 8, 12, 16, 20, 24}。后续校准阶段会进一步确定最终用于评分和 LoRA 编辑的层集合。")
    add_note(
        doc,
        "为什么取最后输入 token。",
        "对所有 token 平均可能稀释指令层面的安全意图；而最后输入 token 是模型即将解码前的条件状态，更直接反映模型如何内部理解图像-指令组合。因此它适合作为风险子空间和 Flow teacher 的表示基础。",
    )

    doc.add_heading("3. 反事实风险子空间构建", level=1)
    add_para(
        doc,
        "风险子空间是本文方法的核心潜变量结构。对于每个 harmful-safe-neighbor 成对样本，在层 l 上计算 harmful 表示与配对 safe-neighbor 表示之间的差分。",
    )
    add_formula(
        doc,
        r"Delta h_{i,l} = h_l(x_i^h) - h_l(x_i^s)",
        r"""
\[
\Delta h_{i,l}=h_l(x_i^h)-h_l(x_i^s)
\]
""",
    )
    add_para(
        doc,
        "该差分用于抵消二者共享的视觉场景、对象语义和任务模板，只保留 harmful trigger 相对于 safe neighbor 额外激活的风险相关偏移。换言之，该方向不是普通语义方向，而是由 paired counterfactual contrast 得到的潜在风险偏移。",
    )
    add_formula(
        doc,
        r"Delta H_l = [Delta h_{1,l}; Delta h_{2,l}; ...; Delta h_{N,l}]",
        r"""
\[
\Delta H_l =
\begin{bmatrix}
\Delta h_{1,l} \\
\Delta h_{2,l} \\
\cdots \\
\Delta h_{N,l}
\end{bmatrix}
\]
""",
    )
    add_para(doc, "将所有配对差分堆叠为矩阵后，对每一层分别进行奇异值分解。")
    add_formula(
        doc,
        r"Delta H_l = U_l Sigma_l V_l^T",
        r"""
\[
\Delta H_l = U_l \Sigma_l V_l^\top
\]
""",
    )
    add_formula(
        doc,
        r"B_l = V_{l,1:k}",
        r"""
\[
B_l = V_{l,1:k}
\]
""",
    )
    add_para(doc, "取前 k 个右奇异向量作为第 l 层风险子空间基 B_l。同时，为 centered score 定义 safe-neighbor center。")
    add_formula(
        doc,
        r"c_l = (1/N) sum_i h_l(x_i^s)",
        r"""
\[
c_l = \frac{1}{N}\sum_{i=1}^{N} h_l(x_i^s)
\]
""",
    )
    add_formula(
        doc,
        r"z_l(x) = (h_l(x)-c_l) B_l^T",
        r"""
\[
z_l(x) = (h_l(x)-c_l)B_l^\top
\]
""",
    )
    add_formula(
        doc,
        r"R_imp^raw(x) = sum_{l in L*} ||z_l(x)||_2",
        r"""
\[
R_{\mathrm{imp}}^{raw}(x)
=
\sum_{l\in\mathcal{L}^\*}
\|z_l(x)\|_2
\]
""",
    )
    add_note(
        doc,
        "如何解释 R_imp。",
        "R_imp 不应被表述为经过校准的通用有害概率，而应被表述为 counterfactual risk-subspace activation，即反事实风险子空间激活强度。retain 样本上的 R_imp 主要用于监测表示漂移，而不是用于跨分布绝对有害性排序。",
    )

    doc.add_heading("4. k 与层选择校准", level=1)
    add_para(
        doc,
        "风险子空间维度 k 和评分模式需要通过验证集校准。核心目标是让风险子空间分数能够区分 harmful trigger 与 paired safe-neighbor，同时避免 k 过大引入过多语义噪声或 retain inflation。",
    )
    add_formula(
        doc,
        r"AUC_{h,s} = AUC(R_imp(x^h), R_imp(x^s))",
        r"""
\[
\mathrm{AUC}_{h,s}
=
\mathrm{AUC}(R_{\mathrm{imp}}(x^h), R_{\mathrm{imp}}(x^s))
\]
""",
    )
    add_formula(
        doc,
        r"Delta_{h,s} = (1/N) sum_i [R_imp(x_i^h)-R_imp(x_i^s)]",
        r"""
\[
\Delta_{h,s}
=
\frac{1}{N}\sum_{i=1}^{N}
\left[
R_{\mathrm{imp}}(x_i^h)-R_{\mathrm{imp}}(x_i^s)
\right]
\]
""",
    )
    add_para(
        doc,
        "实际选择时，先筛选验证集 AUC 接近最优的候选配置，再依据 paired mean difference、retain inflation、更小 k 和 centered score 等规则进行选择。这样可以在可分性和稳定性之间取得折中。",
    )
    add_note(
        doc,
        "为什么需要校准。",
        "较大的 k 可能提高训练集或验证集区分度，但也可能包含与风险无关的场景语义方向。centered score 可以降低部分全局表示偏差。校准过程使风险子空间既有辨别力，又尽量不扩大 retain 样本上的非目标激活。",
    )

    doc.add_heading("5. Risk-transport top-n 编辑层选择", level=1)
    add_para(
        doc,
        "在得到风险子空间后，方法进一步选择少量最适合编辑的层。对于每个候选层，估计 safe-neighbor 到 harmful 的潜在风险迁移强度。每个 pair 的层级迁移向量定义为：",
    )
    add_formula(
        doc,
        r"d_{i,l}=h_l(x_i^h)-h_l(x_i^s)",
        r"""
\[
d_{i,l}=h_l(x_i^h)-h_l(x_i^s)
\]
""",
    )
    add_formula(
        doc,
        r"bar d_l = (1/N) sum_i d_{i,l}",
        r"""
\[
\bar d_l=\frac{1}{N}\sum_i d_{i,l}
\]
""",
    )
    add_formula(
        doc,
        r"RTI_l = (1/N) sum_i ReLU(d_{i,l}^T bar d_l / ||bar d_l||_2)",
        r"""
\[
\mathrm{RTI}_l
=
\frac{1}{N}\sum_i
\mathrm{ReLU}
\left(
d_{i,l}^\top
\frac{\bar d_l}{\|\bar d_l\|_2}
\right)
\]
""",
    )
    add_para(doc, "最终层分数综合风险迁移强度、风险投影比例和 retain overlap 惩罚：")
    add_formula(
        doc,
        r"S_l = RTItilde_l + lambda_p Projtilde_l - lambda_r RetainOverlaptilde_l",
        r"""
\[
S_l
=
\widetilde{\mathrm{RTI}}_l
+
\lambda_p \widetilde{\mathrm{Proj}}_l
-
\lambda_r \widetilde{\mathrm{RetainOverlap}}_l
\]
""",
    )
    add_note(
        doc,
        "为什么只编辑 top-n 层。",
        "全层编辑会增加能力损伤和训练不稳定风险。top-n 选层把 LoRA 容量集中到风险迁移最显著、且与 retain overlap 较低的层，从而更高效地实现安全遗忘。",
    )

    doc.add_heading("6. 基于 ASR 的显性风险评分", level=1)
    add_para(
        doc,
        "显性风险用于衡量模型生成回答中可观察到的 unsafe 行为。当前方法使用安全分类器对生成回答进行判断：若回答包含可执行的危险、违法、自伤或其他有害操作步骤，则显性风险为 1；否则为 0。",
    )
    add_formula(
        doc,
        r"R_exp(x,y)=1 if y is unsafe, otherwise 0",
        r"""
\[
R_{\mathrm{exp}}(x,y)=
\begin{cases}
1, & \text{if } y \text{ is classified as unsafe},\\
0, & \text{otherwise}.
\end{cases}
\]
""",
    )
    add_formula(
        doc,
        r"ASR = (1/|D|) sum_{(x,y) in D} R_exp(x,y)",
        r"""
\[
\mathrm{ASR}
=
\frac{1}{|\mathcal{D}|}
\sum_{(x,y)\in \mathcal{D}}
R_{\mathrm{exp}}(x,y)
\]
""",
    )
    add_note(
        doc,
        "为什么显性与隐性分开。",
        "显性 ASR 反映输出文本是否有害；R_imp 反映模型内部是否仍激活风险子空间。二者分开报告可以避免模型仅通过拒答或文本表面改写掩盖潜在风险，也可以发现潜在风险下降但显性回答仍未完全安全的情况。",
    )

    doc.add_heading("7. 隐性风险归一化与融合总风险", level=1)
    add_para(
        doc,
        "由于 R_imp^raw 是投影范数，其数值范围不固定，因此需要基于训练集分位数进行 min-max 归一化。设训练集 raw score 的下分位数和上分位数分别为 q_alpha 与 q_beta，归一化方式为：",
    )
    add_formula(
        doc,
        r"R_imp(x)=clip((R_imp^raw(x)-q_alpha)/(q_beta-q_alpha),0,1)",
        r"""
\[
R_{\mathrm{imp}}(x)
=
\mathrm{clip}
\left(
\frac{
R_{\mathrm{imp}}^{raw}(x)-q_{\alpha}
}{
q_{\beta}-q_{\alpha}
},
0,1
\right)
\]
""",
    )
    add_para(doc, "融合总风险将显性风险和归一化后的隐性风险进行加权组合。当前主实验采用 0.5/0.5 等权融合。")
    add_formula(
        doc,
        r"R_total(x,y)=lambda_exp R_exp(x,y)+lambda_imp R_imp(x)",
        r"""
\[
R_{\mathrm{total}}(x,y)
=
\lambda_{\mathrm{exp}} R_{\mathrm{exp}}(x,y)
+
\lambda_{\mathrm{imp}} R_{\mathrm{imp}}(x)
\]
""",
    )
    add_formula(
        doc,
        r"R_total = 0.5 R_exp + 0.5 R_imp",
        r"""
\[
R_{\mathrm{total}}
=
0.5R_{\mathrm{exp}}+0.5R_{\mathrm{imp}}
\]
""",
    )
    add_para(doc, "高风险判定采用联合规则，只要显性风险、隐性风险或融合总风险任一超过阈值，即视为高风险。")
    add_formula(
        doc,
        r"I_high = I[R_exp>=0.5 or R_imp>=0.66 or R_total>=0.66]",
        r"""
\[
\mathbb{I}_{high}(x,y)
=
\mathbb{I}
[
R_{\mathrm{exp}}\ge 0.5
\lor
R_{\mathrm{imp}}\ge 0.66
\lor
R_{\mathrm{total}}\ge 0.66
]
\]
""",
    )
    add_note(
        doc,
        "为什么采用 OR 规则。",
        "如果只看 R_total，当 R_exp=0 且 R_imp=1 时，等权融合后 R_total=0.5，可能低于高风险阈值。OR 规则可以避免高隐性风险样本被平均操作掩盖。",
    )

    doc.add_heading("8. Flow Matching Teacher", level=1)
    add_para(
        doc,
        "Flow Matching teacher 用于学习 harmful 隐状态到 safe-neighbor 隐状态之间的平滑潜空间迁移路径。对于每个 pair 和选定层，将 harmful 隐状态作为起点，将 paired safe-neighbor 隐状态作为终点。",
    )
    add_formula(
        doc,
        r"x_0=h_l(x_i^h), x_1=h_l(x_i^s)",
        r"""
\[
x_0=h_l(x_i^h), \quad x_1=h_l(x_i^s)
\]
""",
    )
    add_formula(
        doc,
        r"x_t=(1-t)x_0+t x_1, t~U(0,1)",
        r"""
\[
x_t=(1-t)x_0+t x_1,\quad t\sim U(0,1)
\]
""",
    )
    add_formula(
        doc,
        r"u_t=x_1-x_0",
        r"""
\[
u_t=x_1-x_0
\]
""",
    )
    add_para(
        doc,
        "Flow teacher v_phi 的目标是根据中间状态 x_t、时间 t 和条件 c 预测从 harmful 到 safe-neighbor 的速度方向。条件 c 包含 hidden state、风险子空间系数、显性风险分数和样本类型信息。",
    )
    add_formula(
        doc,
        r"L_velocity=E||v_phi(x_t,t,c)-(x_1-x_0)||_2^2",
        r"""
\[
\mathcal{L}_{velocity}
=
\mathbb{E}_{t,x_0,x_1}
\left[
\|v_\phi(x_t,t,c)-(x_1-x_0)\|_2^2
\right]
\]
""",
    )
    add_formula(
        doc,
        r"L_endpoint=||xhat_1-x_1||_2^2",
        r"""
\[
\mathcal{L}_{endpoint}
=
\|\hat{x}_1-x_1\|_2^2
\]
""",
    )
    add_formula(
        doc,
        r"L_identity=||xhat-x||_2^2",
        r"""
\[
\mathcal{L}_{identity}
=
\|\hat{x}-x\|_2^2
\]
""",
    )
    add_formula(
        doc,
        r"L_flow=L_velocity+lambda_e L_endpoint+lambda_i L_identity",
        r"""
\[
\mathcal{L}_{flow}
=
\mathcal{L}_{velocity}
+
\lambda_e\mathcal{L}_{endpoint}
+
\lambda_i\mathcal{L}_{identity}
\]
""",
    )
    add_note(
        doc,
        "为什么需要 Flow teacher。",
        "单纯文本 CE 可能让模型过度学习固定拒答模板。Flow teacher 提供潜空间中的连续安全迁移方向，使 LoRA 不只是模仿文本答案，而是在隐状态层面朝 safe-neighbor 表示移动。",
    )

    doc.add_heading("9. Flow-guided LoRA Unlearning", level=1)
    add_para(
        doc,
        "Stage 3 冻结基础模型参数，仅训练 LoRA 参数 Delta theta。训练目标由 safe-response 监督、风险子空间对齐、隐性风险压制、安全/保留样本 KL 约束、retain 隐状态保持以及 Flow 蒸馏共同组成。",
    )
    add_formula(
        doc,
        r"L_safeCE=-sum_t log p_{theta+Delta theta}(y_t^s | x^h,y_{<t}^s)",
        r"""
\[
\mathcal{L}_{safeCE}
=
-\sum_t
\log p_{\theta+\Delta\theta}
(y^{s}_t \mid x^h, y^{s}_{<t})
\]
""",
    )
    add_para(doc, "当前主方法中，y^s 来自 paired safe-neighbor response。这样可以引导 harmful prompt 生成安全规避或风险缓解式回答，而不是危险操作步骤。")
    add_formula(
        doc,
        r"z_l^new(x^h)=(h_l^new(x^h)-c_l)B_l^T",
        r"""
\[
z_l^{new}(x^h)=(h_l^{new}(x^h)-c_l)B_l^\top
\]
""",
    )
    add_formula(
        doc,
        r"z_l^old(x^s)=(h_l^old(x^s)-c_l)B_l^T",
        r"""
\[
z_l^{old}(x^s)=(h_l^{old}(x^s)-c_l)B_l^\top
\]
""",
    )
    add_formula(
        doc,
        r"L_align=sum_l ||z_l^new(x^h)-z_l^old(x^s)||_2^2",
        r"""
\[
\mathcal{L}_{align}
=
\sum_{l\in\mathcal{L}^\*}
\left\|
z_l^{new}(x^h)-z_l^{old}(x^s)
\right\|_2^2
\]
""",
    )
    add_formula(
        doc,
        r"L_implicit=sum_l ||z_l^new(x^h)||_2^2",
        r"""
\[
\mathcal{L}_{implicit}
=
\sum_{l\in\mathcal{L}^\*}
\|z_l^{new}(x^h)\|_2^2
\]
""",
    )
    add_para(doc, "L_align 指定 harmful 隐状态应该向哪里移动，L_implicit 直接惩罚残余风险子空间激活。二者分别对应方向约束和风险压制约束。")
    add_formula(
        doc,
        r"L_safeKL=KL(p_theta(.|x^s)||p_{theta+Delta theta}(.|x^s))",
        r"""
\[
\mathcal{L}_{safeKL}
=
\mathrm{KL}
\left(
p_{\theta}(\cdot|x^s)
\|
p_{\theta+\Delta\theta}(\cdot|x^s)
\right)
\]
""",
    )
    add_formula(
        doc,
        r"L_retainKL=KL(p_theta(.|x^r)||p_{theta+Delta theta}(.|x^r))",
        r"""
\[
\mathcal{L}_{retainKL}
=
\mathrm{KL}
\left(
p_{\theta}(\cdot|x^r)
\|
p_{\theta+\Delta\theta}(\cdot|x^r)
\right)
\]
""",
    )
    add_formula(
        doc,
        r"L_retainHidden=sum_l ||h_l^new(x^r)-h_l^old(x^r)||_2^2",
        r"""
\[
\mathcal{L}_{retainHidden}
=
\sum_{l\in\mathcal{L}^\*}
\left\|
h_l^{new}(x^r)-h_l^{old}(x^r)
\right\|_2^2
\]
""",
    )
    add_note(
        doc,
        "为什么要有保持项。",
        "安全遗忘容易变成广泛拒答或能力坍塌。safeKL、retainKL 和 retainHidden 约束更新后模型在安全邻居与保留样本上的行为和表示尽量接近原模型，从而降低副作用。",
    )

    doc.add_heading("10. Flow 蒸馏到 LoRA", level=1)
    add_para(doc, "Flow 蒸馏要求 LoRA 实际产生的隐状态位移与 Flow teacher 给出的安全迁移方向保持一致。LoRA 位移与 Flow 位移分别定义为：")
    add_formula(
        doc,
        r"Delta h_l^LoRA=h_l^new(x^h)-h_l^old(x^h)",
        r"""
\[
\Delta h_l^{LoRA}
=
h_l^{new}(x^h)-h_l^{old}(x^h)
\]
""",
    )
    add_formula(
        doc,
        r"Delta h_l^Flow=xhat_{1,l}^Flow-h_l^old(x^h)",
        r"""
\[
\Delta h_l^{Flow}
=
\hat{x}_{1,l}^{Flow}-h_l^{old}(x^h)
\]
""",
    )
    add_formula(
        doc,
        r"L_flowDelta=||Delta h_l^LoRA-Delta h_l^Flow||_2^2",
        r"""
\[
\mathcal{L}_{flowDelta}
=
\|\Delta h_l^{LoRA}-\Delta h_l^{Flow}\|_2^2
\]
""",
    )
    add_formula(
        doc,
        r"L_flowCos=1-cos(Delta h_l^LoRA, Delta h_l^Flow)",
        r"""
\[
\mathcal{L}_{flowCos}
=
1-
\cos
(
\Delta h_l^{LoRA},
\Delta h_l^{Flow}
)
\]
""",
    )
    add_formula(
        doc,
        r"L_flowRisk=||z_l^LoRA(x^h)-z_l^Flow(x^h)||_2^2",
        r"""
\[
\mathcal{L}_{flowRisk}
=
\left\|
z_l^{LoRA}(x^h)-z_l^{Flow}(x^h)
\right\|_2^2
\]
""",
    )
    add_para(doc, "最终 Stage 3 总损失为所有损失项的加权和：")
    add_formula(
        doc,
        r"L_total=lambda_ce L_safeCE+lambda_align L_align+lambda_imp L_implicit+lambda_safeKL L_safeKL+lambda_retainKL L_retainKL+lambda_hidden L_retainHidden+lambda_flow L_flowDistill",
        r"""
\[
\mathcal{L}_{total}
=
\lambda_{ce}\mathcal{L}_{safeCE}
+
\lambda_{align}\mathcal{L}_{align}
+
\lambda_{imp}\mathcal{L}_{implicit}
+
\lambda_{safeKL}\mathcal{L}_{safeKL}
+
\lambda_{retainKL}\mathcal{L}_{retainKL}
+
\lambda_{hidden}\mathcal{L}_{retainHidden}
+
\lambda_{flow}\mathcal{L}_{flowDistill}
\]
""",
    )
    add_note(
        doc,
        "为什么 lambda_flow 需要 ramp。",
        "训练初期如果 Flow 约束过强，可能干扰 safe CE 和基础行为稳定。采用 ramp 调度可以先建立安全回答倾向，再逐步增强潜空间迁移约束。",
    )

    doc.add_heading("11. 统一严格 before/after 评估", level=1)
    add_para(
        doc,
        "严格评估使用 base model 生成回答作为 before，使用 LoRA model 生成回答作为 after。这样避免了 dataset response 与 generated response 混合比较，使结果真正反映模型更新前后的变化。",
    )
    add_formula(
        doc,
        r"y^base=f_theta(x)",
        r"""
\[
y^{base}=f_{\theta}(x)
\]
""",
    )
    add_formula(
        doc,
        r"y^LoRA=f_{theta+Delta theta}(x)",
        r"""
\[
y^{LoRA}=f_{\theta+\Delta\theta}(x)
\]
""",
    )
    add_para(doc, "对于 harmful trigger，清除率基于 before/after 的组均值计算。")
    add_formula(
        doc,
        r"C_exp=(Rbar_exp^before-Rbar_exp^after)/(Rbar_exp^before+epsilon)",
        r"""
\[
C_{\mathrm{exp}}
=
\frac{
\bar{R}_{\mathrm{exp}}^{before}
-
\bar{R}_{\mathrm{exp}}^{after}
}{
\bar{R}_{\mathrm{exp}}^{before}+\epsilon
}
\]
""",
    )
    add_formula(
        doc,
        r"C_imp=(Rbar_imp^before-Rbar_imp^after)/(Rbar_imp^before+epsilon)",
        r"""
\[
C_{\mathrm{imp}}
=
\frac{
\bar{R}_{\mathrm{imp}}^{before}
-
\bar{R}_{\mathrm{imp}}^{after}
}{
\bar{R}_{\mathrm{imp}}^{before}+\epsilon
}
\]
""",
    )
    add_formula(
        doc,
        r"C_fusion=(Rbar_total^before-Rbar_total^after)/(Rbar_total^before+epsilon)",
        r"""
\[
C_{\mathrm{fusion}}
=
\frac{
\bar{R}_{\mathrm{total}}^{before}
-
\bar{R}_{\mathrm{total}}^{after}
}{
\bar{R}_{\mathrm{total}}^{before}+\epsilon
}
\]
""",
    )
    add_formula(
        doc,
        r"C_balanced=0.5 C_exp+0.5 C_imp",
        r"""
\[
C_{\mathrm{balanced}}
=
0.5C_{\mathrm{exp}}
+
0.5C_{\mathrm{imp}}
\]
""",
    )
    add_formula(
        doc,
        r"RefusalRate=(1/|D|) sum_i I[Refusal(y_i)=1]",
        r"""
\[
\mathrm{RefusalRate}
=
\frac{1}{|\mathcal{D}|}
\sum_i
\mathbb{I}[\mathrm{Refusal}(y_i)=1]
\]
""",
    )
    add_note(
        doc,
        "为什么报告两种总清除率。",
        "融合总风险清除率反映最终风险函数的下降；显隐清除率均衡分则避免单一通道主导结论，要求显性风险和隐性风险都得到改善。",
    )

    doc.add_heading("12. 面向审稿人的解释与局限", level=1)
    add_para(
        doc,
        "审稿人可能质疑 R_imp 只是语义偏差而不是风险分数。对此，论文中应明确：R_imp 不是 calibrated harmfulness probability，而是 paired harmful-safe-neighbor contrast 诱导出的 counterfactual risk-subspace activation。其合理性需要由 harmful-vs-safe-neighbor AUC、paired gap 以及 unlearning 后 harmful 激活下降共同支持。",
    )
    add_formula(
        doc,
        r"R_imp == Counterfactual Risk-Subspace Activation",
        r"""
\[
R_{\mathrm{imp}}
\equiv
\text{Counterfactual Risk-Subspace Activation}
\]
""",
    )
    add_bullets(
        doc,
        [
            "不要把 R_imp 表述为绝对有害概率。",
            "使用 harmful 与 safe-neighbor 的配对可分性验证其风险相关性。",
            "retain 分数主要用于观察方法是否引入表示漂移，而非跨分布比较有害性。",
            "必须同时报告显性 ASR、拒答率和 ROUGE，避免潜在指标掩盖可见行为变化。",
        ],
    )
    add_note(
        doc,
        "建议论文表述。",
        "The implicit score is a counterfactual risk-subspace activation rather than a calibrated harmfulness probability. It is constructed from paired harmful-safe-neighbor differences, which helps suppress shared scene semantics and emphasize unsafe operational deviation.",
    )

    doc.add_section(WD_SECTION.NEW_PAGE)
    doc.add_heading("附录：符号表", level=1)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    widths = [2100, 7260]
    set_table_geometry(table, widths)
    hdr = table.rows[0].cells
    hdr[0].text = "符号"
    hdr[1].text = "含义"
    for cell in hdr:
        set_cell_shading(cell, "F4F6F9")
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
    entries = [
        ("x_i^h", "harmful trigger 样本。"),
        ("x_i^s", "配对 safe-neighbor 样本。"),
        ("x_i^r", "retain 样本。"),
        ("h_l(x)", "第 l 层最后输入 token 隐状态。"),
        ("B_l", "第 l 层 top-k 风险子空间基。"),
        ("c_l", "第 l 层 safe-neighbor center。"),
        ("R_exp", "由安全分类器得到的显性 ASR 风险。"),
        ("R_imp", "归一化后的反事实风险子空间激活。"),
        ("R_total", "显性与隐性风险的加权融合。"),
        ("C_fusion", "基于融合总风险的清除率。"),
        ("C_balanced", "显性清除率与隐性清除率的均衡分。"),
    ]
    for sym, meaning in entries:
        cells = table.add_row().cells
        cells[0].text = sym
        cells[1].text = meaning
        set_cell_width(cells[0], widths[0])
        set_cell_width(cells[1], widths[1])

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
