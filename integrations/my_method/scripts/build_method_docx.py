from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("all.docx").resolve()


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def shade_paragraph(paragraph, fill: str = "F4F6F9") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int]) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths_dxa[idx])


def add_field(paragraph, field: str) -> None:
    run = paragraph.add_run()
    fld_char_1 = OxmlElement("w:fldChar")
    fld_char_1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = field
    fld_char_2 = OxmlElement("w:fldChar")
    fld_char_2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_1)
    run._r.append(instr_text)
    run._r.append(fld_char_2)


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor(0, 0, 0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    title = styles["Title"]
    title.font.name = "Calibri"
    title._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    title.font.size = Pt(22)
    title.font.color.rgb = RGBColor(11, 37, 69)
    title.paragraph_format.space_after = Pt(6)

    subtitle = styles["Subtitle"]
    subtitle.font.name = "Calibri"
    subtitle._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    subtitle.font.size = Pt(11)
    subtitle.font.color.rgb = RGBColor(85, 85, 85)

    for name, size, color, before, after in [
        ("Heading 1", 16, RGBColor(46, 116, 181), 18, 10),
        ("Heading 2", 13, RGBColor(46, 116, 181), 12, 6),
        ("Heading 3", 12, RGBColor(31, 77, 120), 8, 4),
    ]:
        st = styles[name]
        st.font.name = "Calibri"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        st.font.size = Pt(size)
        st.font.color.rgb = color
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True

    code = styles.add_style("CodeBlock", 1)
    code.font.name = "Consolas"
    code._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    code.font.size = Pt(9)
    code.font.color.rgb = RGBColor(31, 58, 95)
    code.paragraph_format.left_indent = Inches(0.18)
    code.paragraph_format.right_indent = Inches(0.18)
    code.paragraph_format.space_before = Pt(3)
    code.paragraph_format.space_after = Pt(8)
    code.paragraph_format.line_spacing = 1.08

    formula = styles.add_style("Formula", 1)
    formula.font.name = "Cambria Math"
    formula._element.rPr.rFonts.set(qn("w:eastAsia"), "Cambria Math")
    formula.font.size = Pt(10.5)
    formula.font.color.rgb = RGBColor(0, 0, 0)
    formula.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    formula.paragraph_format.space_before = Pt(3)
    formula.paragraph_format.space_after = Pt(3)

    note = styles.add_style("NoteText", 1)
    note.font.name = "Calibri"
    note._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    note.font.size = Pt(10)
    note.font.color.rgb = RGBColor(31, 58, 95)
    note.paragraph_format.left_indent = Inches(0.16)
    note.paragraph_format.right_indent = Inches(0.16)
    note.paragraph_format.space_before = Pt(4)
    note.paragraph_format.space_after = Pt(8)
    note.paragraph_format.line_spacing = 1.25

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.add_run("Risk-subspace Flow-guided LoRA Method | Page ")
    add_field(footer, "PAGE")


def add_para(doc: Document, text: str, style: str | None = None, bold_prefix: str | None = None) -> None:
    p = doc.add_paragraph(style=style)
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        r.bold = True
        p.add_run(text[len(bold_prefix):])
    else:
        p.add_run(text)


def add_formula(doc: Document, display: str, latex: str) -> None:
    p = doc.add_paragraph(display, style="Formula")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    label = doc.add_paragraph("LaTeX code:", style="NoteText")
    label.runs[0].bold = True
    code = doc.add_paragraph(latex.strip(), style="CodeBlock")
    shade_paragraph(code)


def add_note(doc: Document, title: str, text: str) -> None:
    p = doc.add_paragraph(style="NoteText")
    shade_paragraph(p)
    r = p.add_run(title + " ")
    r.bold = True
    p.add_run(text)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.375)
        p.paragraph_format.first_line_indent = Inches(-0.194)
        p.paragraph_format.space_after = Pt(4)
        p.add_run(item)


def add_summary_table(doc: Document) -> None:
    doc.add_heading("Method components and their roles", level=2)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    widths = [1900, 3950, 3510]
    set_table_geometry(table, widths)
    hdr = table.rows[0].cells
    for i, text in enumerate(["Component", "Role in the method", "Why it is useful"]):
        hdr[i].text = text
        set_cell_shading(hdr[i], "F4F6F9")
        for p in hdr[i].paragraphs:
            for run in p.runs:
                run.bold = True
    rows = [
        (
            "Risk subspace",
            "Extracts latent directions from paired harmful-safe-neighbor representation differences.",
            "Controls for shared scene semantics and focuses on counterfactual unsafe operational deviation.",
        ),
        (
            "Risk-transport layer selection",
            "Selects layers whose safe-to-harmful movement is most aligned with the risk subspace.",
            "Limits LoRA edits to a small number of high-impact layers, reducing collateral damage.",
        ),
        (
            "Flow Matching teacher",
            "Learns a continuous hidden-space path from harmful states toward safe-neighbor states.",
            "Provides a smooth target direction instead of relying only on hard CE supervision.",
        ),
        (
            "Flow-guided LoRA",
            "Trains lightweight adapters using safe response CE, risk projection losses, KL retention, and flow distillation.",
            "Suppresses unsafe latent behavior while preserving the base model and retain-set behavior.",
        ),
        (
            "Unified evaluation",
            "Reports explicit ASR, implicit risk-subspace activation, fused total risk, refusal rate, ROUGE, and clearance scores.",
            "Separates observable unsafe behavior from latent risk activation and capability preservation.",
        ),
    ]
    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            cells[i].text = text
            set_cell_width(cells[i], widths[i])
    doc.add_paragraph()


def main() -> None:
    doc = Document()
    style_document(doc)

    doc.add_paragraph("Risk-Subspace Flow-guided LoRA Unlearning", style="Title")
    doc.add_paragraph("A detailed method description with formulas and reusable LaTeX code", style="Subtitle")
    add_note(
        doc,
        "Scope.",
        "This document describes the current main method: SafeNb-based counterfactual risk-subspace construction, risk-transport top-n layer selection, Flow Matching teacher, Flow-guided LoRA unlearning, and unified strict before/after evaluation. It intentionally excludes the withdrawn harmful-only variant.",
    )

    add_summary_table(doc)

    doc.add_heading("1. Problem setup and data organization", level=1)
    add_para(
        doc,
        "The goal is to reduce a multimodal model's tendency to produce unsafe operational responses for harmful triggers, while preserving behavior on safe-neighbor and retain samples. The model is denoted by f_theta. Each input consists of an image v and a text instruction x. The dataset is organized into paired or grouped examples containing a harmful trigger x_i^h, a safe neighbor x_i^s, and a retain sample x_i^r.",
    )
    add_formula(
        doc,
        r"D = {(x_i^h, x_i^s, x_i^r)}_{i=1}^N",
        r"""
\[
\mathcal{D}=\{(x_i^h, x_i^s, x_i^r)\}_{i=1}^{N}
\]
""",
    )
    add_para(
        doc,
        "Here, x_i^h denotes a harmful trigger; x_i^s denotes the paired safe-neighbor sample; x_i^r denotes a retain sample. The safe neighbor is not a random safe example. It is intentionally close to the harmful trigger in visual or task context but differs in the intended safe response. This paired construction is central to reducing semantic confounding.",
    )
    add_note(
        doc,
        "Why this design.",
        "If harmful examples were compared only with unrelated retain samples, the learned direction could mainly capture topic or scene differences. Pairing harmful triggers with safe-neighbor controls encourages the method to isolate the unsafe operational component rather than the broad visual context.",
    )

    doc.add_heading("2. Hidden-state representation", level=1)
    add_para(
        doc,
        "For each multimodal input, the method extracts hidden states from selected model layers. In the current implementation, the representation is the last input token hidden state at layer l. This token is used because it is the model state immediately before generation and has already integrated the image content and the full text instruction.",
    )
    add_formula(
        doc,
        r"h_l(x) = LastTokenHidden_l(f_theta(v,x))",
        r"""
\[
h_l(x)=\mathrm{LastTokenHidden}_l(f_\theta(v,x))
\]
""",
    )
    add_para(
        doc,
        "The layer index l belongs to a candidate layer set L, for example {4, 8, 12, 16, 20, 24}. Later stages calibrate which subset of layers should be used for scoring and LoRA editing.",
    )
    add_note(
        doc,
        "Why last-token hidden states.",
        "Averaging over all tokens can blur instruction-specific safety intent, whereas the last input token is the conditional state used to start decoding. It is therefore a compact representation of how the model has internally interpreted the image-instruction pair before producing an answer.",
    )

    doc.add_heading("3. Counterfactual risk-subspace construction", level=1)
    add_para(
        doc,
        "The central latent object is a counterfactual risk subspace. For each paired harmful-safe-neighbor example, the method computes the difference between the harmful representation and the paired safe-neighbor representation at the same layer.",
    )
    add_formula(
        doc,
        r"Delta h_{i,l} = h_l(x_i^h) - h_l(x_i^s)",
        r"""
\[
\Delta h_{i,l}=h_l(x_i^h)-h_l(x_i^s)
\]
""",
    )
    add_para(
        doc,
        "The difference vector is intended to remove shared scene-level information. For example, both samples may involve the same dangerous visual context, but only the harmful trigger encourages unsafe operational execution. The residual direction is therefore treated as a latent unsafe-deviation signal.",
    )
    add_formula(
        doc,
        r"Delta H_l = [Delta h_{1,l}; Delta h_{2,l}; ...; Delta h_{N,l}]",
        r"""
\[
\Delta H_l =
\begin{bmatrix}
\Delta h_{1,l} \\
\Delta h_{2,l} \\
\cdots \\
\Delta h_{N,l}
\end{bmatrix}
\]
""",
    )
    add_para(doc, "For each layer, the paired differences are stacked into a matrix and decomposed with singular value decomposition.")
    add_formula(
        doc,
        r"Delta H_l = U_l Sigma_l V_l^T",
        r"""
\[
\Delta H_l = U_l \Sigma_l V_l^\top
\]
""",
    )
    add_formula(
        doc,
        r"B_l = V_{l,1:k}",
        r"""
\[
B_l = V_{l,1:k}
\]
""",
    )
    add_para(
        doc,
        "The first k right singular vectors form the basis B_l of the risk subspace. The method also computes a safe-neighbor center for centered projection scoring.",
    )
    add_formula(
        doc,
        r"c_l = (1/N) sum_i h_l(x_i^s)",
        r"""
\[
c_l = \frac{1}{N}\sum_{i=1}^{N} h_l(x_i^s)
\]
""",
    )
    add_formula(
        doc,
        r"z_l(x) = (h_l(x)-c_l) B_l^T",
        r"""
\[
z_l(x) = (h_l(x)-c_l)B_l^\top
\]
""",
    )
    add_formula(
        doc,
        r"R_imp^raw(x) = sum_{l in L*} ||z_l(x)||_2",
        r"""
\[
R_{\mathrm{imp}}^{raw}(x)
=
\sum_{l\in\mathcal{L}^\*}
\|z_l(x)\|_2
\]
""",
    )
    add_note(
        doc,
        "Interpretation.",
        "R_imp is best described as counterfactual risk-subspace activation, not as a calibrated universal harmfulness probability. It measures how strongly a representation activates directions learned from harmful-safe-neighbor contrasts. Retain scores should therefore be interpreted mainly as representation drift diagnostics, not as cross-domain absolute harmfulness ranks.",
    )

    doc.add_heading("4. Calibration of k and layer selection", level=1)
    add_para(
        doc,
        "The dimension k and the risk-scoring mode are selected by validation performance. The method evaluates whether the risk-subspace score separates harmful triggers from safe-neighbor controls and whether the paired harmful score is consistently larger than the paired safe-neighbor score before unlearning.",
    )
    add_formula(
        doc,
        r"AUC_{h,s} = AUC(R_imp(x^h), R_imp(x^s))",
        r"""
\[
\mathrm{AUC}_{h,s}
=
\mathrm{AUC}(R_{\mathrm{imp}}(x^h), R_{\mathrm{imp}}(x^s))
\]
""",
    )
    add_formula(
        doc,
        r"Delta_{h,s} = (1/N) sum_i [R_imp(x_i^h)-R_imp(x_i^s)]",
        r"""
\[
\Delta_{h,s}
=
\frac{1}{N}\sum_{i=1}^{N}
\left[
R_{\mathrm{imp}}(x_i^h)-R_{\mathrm{imp}}(x_i^s)
\right]
\]
""",
    )
    add_para(
        doc,
        "In practice, configurations within a small tolerance of the best validation AUC are compared using paired mean difference, retain inflation, smaller k, and centered score preference. This avoids selecting a high-dimensional subspace that overfits or inflates retain scores.",
    )
    add_note(
        doc,
        "Why calibration is necessary.",
        "A larger k can improve separability but may include nuisance directions. A centered score can reduce some global representation bias. Calibration balances harmful-safe separability against stability on retain data.",
    )

    doc.add_heading("5. Risk-transport top-n layer selection", level=1)
    add_para(
        doc,
        "After identifying the risk subspace, the method selects a small set of layers for LoRA editing. For each candidate layer, it estimates how strongly safe-neighbor-to-harmful movement aligns with a consistent transport direction. The per-pair transport vector is defined as follows.",
    )
    add_formula(
        doc,
        r"d_{i,l}=h_l(x_i^h)-h_l(x_i^s)",
        r"""
\[
d_{i,l}=h_l(x_i^h)-h_l(x_i^s)
\]
""",
    )
    add_formula(
        doc,
        r"bar d_l = (1/N) sum_i d_{i,l}",
        r"""
\[
\bar d_l=\frac{1}{N}\sum_i d_{i,l}
\]
""",
    )
    add_formula(
        doc,
        r"RTI_l = (1/N) sum_i ReLU(d_{i,l}^T bar d_l / ||bar d_l||_2)",
        r"""
\[
\mathrm{RTI}_l
=
\frac{1}{N}\sum_i
\mathrm{ReLU}
\left(
d_{i,l}^\top
\frac{\bar d_l}{\|\bar d_l\|_2}
\right)
\]
""",
    )
    add_para(
        doc,
        "The final layer score combines normalized risk-transport influence, risk projection ratio, and a retain-overlap penalty.",
    )
    add_formula(
        doc,
        r"S_l = RTItilde_l + lambda_p Projtilde_l - lambda_r RetainOverlaptilde_l",
        r"""
\[
S_l
=
\widetilde{\mathrm{RTI}}_l
+
\lambda_p \widetilde{\mathrm{Proj}}_l
-
\lambda_r \widetilde{\mathrm{RetainOverlap}}_l
\]
""",
    )
    add_note(
        doc,
        "Why top-n layer selection.",
        "Editing all layers is unnecessary and can damage general capability. The top-n criterion concentrates LoRA capacity on layers where the harmful-safe transport is strongest while penalizing layers that overlap heavily with retain representations.",
    )

    doc.add_heading("6. Explicit risk scoring with ASR", level=1)
    add_para(
        doc,
        "Visible unsafe behavior is measured using an ASR-style explicit risk score. A safety classifier evaluates the model response. If the response provides unsafe operational, procedural, or step-by-step content, the explicit score is one; otherwise it is zero.",
    )
    add_formula(
        doc,
        r"R_exp(x,y)=1 if y is unsafe, otherwise 0",
        r"""
\[
R_{\mathrm{exp}}(x,y)=
\begin{cases}
1, & \text{if } y \text{ is classified as unsafe},\\
0, & \text{otherwise}.
\end{cases}
\]
""",
    )
    add_formula(
        doc,
        r"ASR = (1/|D|) sum_{(x,y) in D} R_exp(x,y)",
        r"""
\[
\mathrm{ASR}
=
\frac{1}{|\mathcal{D}|}
\sum_{(x,y)\in \mathcal{D}}
R_{\mathrm{exp}}(x,y)
\]
""",
    )
    add_note(
        doc,
        "Why ASR is separated from R_imp.",
        "Explicit ASR captures observable unsafe responses, whereas R_imp captures latent risk-subspace activation. Reporting both prevents a model from appearing safe only because it reduces one channel while worsening the other.",
    )

    doc.add_heading("7. Normalization and fused total risk", level=1)
    add_para(
        doc,
        "The raw implicit score is a projection norm and therefore not naturally bounded. The method normalizes it using percentile min-max statistics computed on the training set. The current configuration uses the lower percentile q_alpha and upper percentile q_beta, with clipping to [0,1].",
    )
    add_formula(
        doc,
        r"R_imp(x)=clip((R_imp^raw(x)-q_alpha)/(q_beta-q_alpha),0,1)",
        r"""
\[
R_{\mathrm{imp}}(x)
=
\mathrm{clip}
\left(
\frac{
R_{\mathrm{imp}}^{raw}(x)-q_{\alpha}
}{
q_{\beta}-q_{\alpha}
},
0,1
\right)
\]
""",
    )
    add_para(
        doc,
        "The fused total risk combines explicit ASR and normalized implicit activation using weights lambda_exp and lambda_imp. In the current evaluation, both weights are 0.5.",
    )
    add_formula(
        doc,
        r"R_total(x,y)=lambda_exp R_exp(x,y)+lambda_imp R_imp(x)",
        r"""
\[
R_{\mathrm{total}}(x,y)
=
\lambda_{\mathrm{exp}} R_{\mathrm{exp}}(x,y)
+
\lambda_{\mathrm{imp}} R_{\mathrm{imp}}(x)
\]
""",
    )
    add_formula(
        doc,
        r"R_total = 0.5 R_exp + 0.5 R_imp",
        r"""
\[
R_{\mathrm{total}}
=
0.5R_{\mathrm{exp}}+0.5R_{\mathrm{imp}}
\]
""",
    )
    add_para(
        doc,
        "The high-risk decision is not based only on R_total. The current evaluation marks a sample as high risk if any of the explicit, implicit, or fused risk thresholds is exceeded.",
    )
    add_formula(
        doc,
        r"I_high = I[R_exp>=0.5 or R_imp>=0.66 or R_total>=0.66]",
        r"""
\[
\mathbb{I}_{high}(x,y)
=
\mathbb{I}
[
R_{\mathrm{exp}}\ge 0.5
\lor
R_{\mathrm{imp}}\ge 0.66
\lor
R_{\mathrm{total}}\ge 0.66
]
\]
""",
    )
    add_note(
        doc,
        "Why the OR rule.",
        "With equal fusion, a sample can have R_exp=0 and R_imp=1 but R_total=0.5. The OR rule prevents such samples from being hidden by averaging and therefore gives a stricter high-risk accounting.",
    )

    doc.add_heading("8. Flow Matching teacher", level=1)
    add_para(
        doc,
        "The Flow Matching teacher learns a smooth hidden-space transition from harmful states to safe-neighbor states. For each pair and selected layer, the harmful hidden state is the source x_0 and the paired safe-neighbor hidden state is the target x_1.",
    )
    add_formula(
        doc,
        r"x_0=h_l(x_i^h), x_1=h_l(x_i^s)",
        r"""
\[
x_0=h_l(x_i^h), \quad x_1=h_l(x_i^s)
\]
""",
    )
    add_formula(
        doc,
        r"x_t=(1-t)x_0+t x_1, t~U(0,1)",
        r"""
\[
x_t=(1-t)x_0+t x_1,\quad t\sim U(0,1)
\]
""",
    )
    add_formula(
        doc,
        r"u_t=x_1-x_0",
        r"""
\[
u_t=x_1-x_0
\]
""",
    )
    add_para(
        doc,
        "The teacher v_phi predicts the velocity needed to move from the harmful hidden state toward the safe-neighbor target, conditioned on the hidden state, risk-subspace coefficients, explicit risk, time, and group metadata.",
    )
    add_formula(
        doc,
        r"L_velocity=E||v_phi(x_t,t,c)-(x_1-x_0)||_2^2",
        r"""
\[
\mathcal{L}_{velocity}
=
\mathbb{E}_{t,x_0,x_1}
\left[
\|v_\phi(x_t,t,c)-(x_1-x_0)\|_2^2
\right]
\]
""",
    )
    add_formula(
        doc,
        r"L_endpoint=||xhat_1-x_1||_2^2",
        r"""
\[
\mathcal{L}_{endpoint}
=
\|\hat{x}_1-x_1\|_2^2
\]
""",
    )
    add_formula(
        doc,
        r"L_identity=||xhat-x||_2^2",
        r"""
\[
\mathcal{L}_{identity}
=
\|\hat{x}-x\|_2^2
\]
""",
    )
    add_formula(
        doc,
        r"L_flow=L_velocity+lambda_e L_endpoint+lambda_i L_identity",
        r"""
\[
\mathcal{L}_{flow}
=
\mathcal{L}_{velocity}
+
\lambda_e\mathcal{L}_{endpoint}
+
\lambda_i\mathcal{L}_{identity}
\]
""",
    )
    add_note(
        doc,
        "Why a Flow teacher.",
        "A direct CE target can force a narrow refusal style or overfit to templates. The Flow teacher supplies a continuous latent transport direction, helping LoRA move harmful states toward safe-neighbor states in representation space rather than relying only on text-level supervision.",
    )

    doc.add_heading("9. Flow-guided LoRA unlearning", level=1)
    add_para(
        doc,
        "The base model is frozen and only LoRA parameters Delta theta are optimized. The training objective combines safe-response supervision, risk-subspace alignment, implicit risk suppression, KL preservation, retain hidden-state preservation, and Flow teacher distillation.",
    )
    add_formula(
        doc,
        r"L_safeCE=-sum_t log p_{theta+Delta theta}(y_t^s | x^h,y_{<t}^s)",
        r"""
\[
\mathcal{L}_{safeCE}
=
-\sum_t
\log p_{\theta+\Delta\theta}
(y^{s}_t \mid x^h, y^{s}_{<t})
\]
""",
    )
    add_para(
        doc,
        "In the current main method, y^s is the paired safe-neighbor response. This encourages the model to answer harmful prompts with safe or risk-mitigating content rather than unsafe operational instructions.",
    )
    add_formula(
        doc,
        r"z_l^new(x^h)=(h_l^new(x^h)-c_l)B_l^T",
        r"""
\[
z_l^{new}(x^h)=(h_l^{new}(x^h)-c_l)B_l^\top
\]
""",
    )
    add_formula(
        doc,
        r"z_l^old(x^s)=(h_l^old(x^s)-c_l)B_l^T",
        r"""
\[
z_l^{old}(x^s)=(h_l^{old}(x^s)-c_l)B_l^\top
\]
""",
    )
    add_formula(
        doc,
        r"L_align=sum_l ||z_l^new(x^h)-z_l^old(x^s)||_2^2",
        r"""
\[
\mathcal{L}_{align}
=
\sum_{l\in\mathcal{L}^\*}
\left\|
z_l^{new}(x^h)-z_l^{old}(x^s)
\right\|_2^2
\]
""",
    )
    add_formula(
        doc,
        r"L_implicit=sum_l ||z_l^new(x^h)||_2^2",
        r"""
\[
\mathcal{L}_{implicit}
=
\sum_{l\in\mathcal{L}^\*}
\|z_l^{new}(x^h)\|_2^2
\]
""",
    )
    add_para(
        doc,
        "The alignment term pulls harmful prompt representations toward the safe-neighbor projection, while the implicit term directly suppresses risk-subspace activation. These two losses are complementary: one specifies where to move, and the other penalizes residual risk activation.",
    )
    add_formula(
        doc,
        r"L_safeKL=KL(p_theta(.|x^s)||p_{theta+Delta theta}(.|x^s))",
        r"""
\[
\mathcal{L}_{safeKL}
=
\mathrm{KL}
\left(
p_{\theta}(\cdot|x^s)
\|
p_{\theta+\Delta\theta}(\cdot|x^s)
\right)
\]
""",
    )
    add_formula(
        doc,
        r"L_retainKL=KL(p_theta(.|x^r)||p_{theta+Delta theta}(.|x^r))",
        r"""
\[
\mathcal{L}_{retainKL}
=
\mathrm{KL}
\left(
p_{\theta}(\cdot|x^r)
\|
p_{\theta+\Delta\theta}(\cdot|x^r)
\right)
\]
""",
    )
    add_formula(
        doc,
        r"L_retainHidden=sum_l ||h_l^new(x^r)-h_l^old(x^r)||_2^2",
        r"""
\[
\mathcal{L}_{retainHidden}
=
\sum_{l\in\mathcal{L}^\*}
\left\|
h_l^{new}(x^r)-h_l^{old}(x^r)
\right\|_2^2
\]
""",
    )
    add_note(
        doc,
        "Why preservation losses.",
        "Unlearning can easily become broad refusal or capability damage. KL and hidden-state preservation constrain the updated model to remain close to the base model on safe-neighbor and retain inputs.",
    )

    doc.add_heading("10. Flow distillation into LoRA", level=1)
    add_para(
        doc,
        "Flow distillation makes the LoRA-induced hidden-state movement agree with the Flow teacher's safe transport direction. The LoRA displacement and Flow displacement are defined below.",
    )
    add_formula(
        doc,
        r"Delta h_l^LoRA=h_l^new(x^h)-h_l^old(x^h)",
        r"""
\[
\Delta h_l^{LoRA}
=
h_l^{new}(x^h)-h_l^{old}(x^h)
\]
""",
    )
    add_formula(
        doc,
        r"Delta h_l^Flow=xhat_{1,l}^Flow-h_l^old(x^h)",
        r"""
\[
\Delta h_l^{Flow}
=
\hat{x}_{1,l}^{Flow}-h_l^{old}(x^h)
\]
""",
    )
    add_formula(
        doc,
        r"L_flowDelta=||Delta h_l^LoRA-Delta h_l^Flow||_2^2",
        r"""
\[
\mathcal{L}_{flowDelta}
=
\|\Delta h_l^{LoRA}-\Delta h_l^{Flow}\|_2^2
\]
""",
    )
    add_formula(
        doc,
        r"L_flowCos=1-cos(Delta h_l^LoRA, Delta h_l^Flow)",
        r"""
\[
\mathcal{L}_{flowCos}
=
1-
\cos
(
\Delta h_l^{LoRA},
\Delta h_l^{Flow}
)
\]
""",
    )
    add_formula(
        doc,
        r"L_flowRisk=||z_l^LoRA(x^h)-z_l^Flow(x^h)||_2^2",
        r"""
\[
\mathcal{L}_{flowRisk}
=
\left\|
z_l^{LoRA}(x^h)-z_l^{Flow}(x^h)
\right\|_2^2
\]
""",
    )
    add_para(doc, "The final Stage 3 objective is a weighted sum of all losses.")
    add_formula(
        doc,
        r"L_total=lambda_ce L_safeCE+lambda_align L_align+lambda_imp L_implicit+lambda_safeKL L_safeKL+lambda_retainKL L_retainKL+lambda_hidden L_retainHidden+lambda_flow L_flowDistill",
        r"""
\[
\mathcal{L}_{total}
=
\lambda_{ce}\mathcal{L}_{safeCE}
+
\lambda_{align}\mathcal{L}_{align}
+
\lambda_{imp}\mathcal{L}_{implicit}
+
\lambda_{safeKL}\mathcal{L}_{safeKL}
+
\lambda_{retainKL}\mathcal{L}_{retainKL}
+
\lambda_{hidden}\mathcal{L}_{retainHidden}
+
\lambda_{flow}\mathcal{L}_{flowDistill}
\]
""",
    )
    add_note(
        doc,
        "Why ramping lambda_flow.",
        "The Flow term is ramped rather than applied at full strength from the first step. Early training first stabilizes safe-response behavior; later training increasingly follows the teacher's latent transport direction.",
    )

    doc.add_heading("11. Unified strict before/after evaluation", level=1)
    add_para(
        doc,
        "The strict evaluation uses generated base-model responses as before and generated LoRA-model responses as after. This avoids comparing dataset answers with generated answers and gives a model-to-model before/after comparison.",
    )
    add_formula(
        doc,
        r"y^base=f_theta(x)",
        r"""
\[
y^{base}=f_{\theta}(x)
\]
""",
    )
    add_formula(
        doc,
        r"y^LoRA=f_{theta+Delta theta}(x)",
        r"""
\[
y^{LoRA}=f_{\theta+\Delta\theta}(x)
\]
""",
    )
    add_para(doc, "For harmful triggers, clearance rates are computed from before and after group means.")
    add_formula(
        doc,
        r"C_exp=(Rbar_exp^before-Rbar_exp^after)/(Rbar_exp^before+epsilon)",
        r"""
\[
C_{\mathrm{exp}}
=
\frac{
\bar{R}_{\mathrm{exp}}^{before}
-
\bar{R}_{\mathrm{exp}}^{after}
}{
\bar{R}_{\mathrm{exp}}^{before}+\epsilon
}
\]
""",
    )
    add_formula(
        doc,
        r"C_imp=(Rbar_imp^before-Rbar_imp^after)/(Rbar_imp^before+epsilon)",
        r"""
\[
C_{\mathrm{imp}}
=
\frac{
\bar{R}_{\mathrm{imp}}^{before}
-
\bar{R}_{\mathrm{imp}}^{after}
}{
\bar{R}_{\mathrm{imp}}^{before}+\epsilon
}
\]
""",
    )
    add_formula(
        doc,
        r"C_fusion=(Rbar_total^before-Rbar_total^after)/(Rbar_total^before+epsilon)",
        r"""
\[
C_{\mathrm{fusion}}
=
\frac{
\bar{R}_{\mathrm{total}}^{before}
-
\bar{R}_{\mathrm{total}}^{after}
}{
\bar{R}_{\mathrm{total}}^{before}+\epsilon
}
\]
""",
    )
    add_formula(
        doc,
        r"C_balanced=0.5 C_exp+0.5 C_imp",
        r"""
\[
C_{\mathrm{balanced}}
=
0.5C_{\mathrm{exp}}
+
0.5C_{\mathrm{imp}}
\]
""",
    )
    add_formula(
        doc,
        r"RefusalRate=(1/|D|) sum_i I[Refusal(y_i)=1]",
        r"""
\[
\mathrm{RefusalRate}
=
\frac{1}{|\mathcal{D}|}
\sum_i
\mathbb{I}[\mathrm{Refusal}(y_i)=1]
\]
""",
    )
    add_note(
        doc,
        "Why two total-clearance scores.",
        "Fusion total risk clearance measures the decrease in the final fused risk score. Balanced clearance prevents one channel from dominating the conclusion by averaging explicit and implicit clearance rates directly.",
    )

    doc.add_heading("12. Reviewer-facing interpretation and caveats", level=1)
    add_para(
        doc,
        "A potential reviewer concern is that R_imp may capture semantic deviation rather than risk. The method addresses this by defining R_imp as a counterfactual risk-subspace activation rather than a calibrated harmfulness probability. Its validity should be supported empirically by harmful-vs-safe-neighbor AUC, paired harmful-safe gaps, and before/after suppression on harmful triggers.",
    )
    add_formula(
        doc,
        r"R_imp == Counterfactual Risk-Subspace Activation",
        r"""
\[
R_{\mathrm{imp}}
\equiv
\text{Counterfactual Risk-Subspace Activation}
\]
""",
    )
    add_bullets(
        doc,
        [
            "Do not claim that R_imp is an absolute harmfulness probability.",
            "Use harmful and safe-neighbor scores to validate risk-related separability under paired controls.",
            "Use retain scores primarily as drift diagnostics, because retain samples come from a broader distribution.",
            "Report explicit ASR and refusal rate alongside R_imp so observable safety behavior remains visible.",
        ],
    )
    add_note(
        doc,
        "Recommended wording.",
        "The implicit score measures activation magnitude in a task-specific counterfactual risk subspace. Because the subspace is constructed from paired harmful-safe-neighbor contrasts, it suppresses part of the shared scene semantics. However, it should be interpreted as a latent safety signal rather than a universal risk probability.",
    )

    doc.add_section(WD_SECTION.NEW_PAGE)
    doc.add_heading("Appendix: compact notation table", level=1)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    widths = [2100, 7260]
    set_table_geometry(table, widths)
    hdr = table.rows[0].cells
    hdr[0].text = "Symbol"
    hdr[1].text = "Meaning"
    for cell in hdr:
        set_cell_shading(cell, "F4F6F9")
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
    entries = [
        ("x_i^h", "Harmful trigger sample."),
        ("x_i^s", "Paired safe-neighbor sample."),
        ("x_i^r", "Retain sample."),
        ("h_l(x)", "Last input token hidden state at layer l."),
        ("B_l", "Top-k risk-subspace basis at layer l."),
        ("c_l", "Safe-neighbor center at layer l."),
        ("R_exp", "Explicit ASR score from safety classification."),
        ("R_imp", "Normalized counterfactual risk-subspace activation."),
        ("R_total", "Weighted fusion of explicit and implicit risk."),
        ("C_fusion", "Clearance rate computed on fused total risk."),
        ("C_balanced", "Average of explicit and implicit clearance rates."),
    ]
    for sym, meaning in entries:
        cells = table.add_row().cells
        cells[0].text = sym
        cells[1].text = meaning
        set_cell_width(cells[0], widths[0])
        set_cell_width(cells[1], widths[1])

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
